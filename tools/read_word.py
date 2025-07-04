import pathlib

from docx import Document
from generative_ai_toolkit.agent import registry

BASE_DIR = pathlib.Path.cwd().resolve()


def _resolve_path(path: str) -> pathlib.Path:
    """
    Resolve a path safely relative to the base directory and prevent path traversal.

    Parameters
    -----
    path : str
        The relative or absolute path to resolve
    """
    abs_path = (BASE_DIR / path).resolve()
    if not str(abs_path).startswith(str(BASE_DIR)):
        raise PermissionError("Access outside of the working directory is not allowed")
    return abs_path


def _format_paragraph_with_style(paragraph) -> str:
    """
    Format a paragraph based on its style, converting Word heading styles to markdown.
    
    Parameters
    -----
    paragraph : docx.text.paragraph.Paragraph
        The paragraph object from python-docx
        
    Returns
    -----
    str
        The formatted text with appropriate markdown headings
    """
    text = paragraph.text.strip()
    if not text:
        return ""
    
    style_name = paragraph.style.name.lower()
    
    # Convert heading styles to markdown
    if 'heading 1' in style_name:
        return f"# {text}"
    elif 'heading 2' in style_name:
        return f"## {text}"
    elif 'heading 3' in style_name:
        return f"### {text}"
    elif 'heading 4' in style_name:
        return f"#### {text}"
    elif 'heading 5' in style_name:
        return f"##### {text}"
    elif 'heading 6' in style_name:
        return f"###### {text}"
    elif 'title' in style_name:
        return f"# {text}"  # Treat title as main heading
    elif 'subtitle' in style_name:
        return f"## {text}"  # Treat subtitle as level 2 heading
    else:
        # Regular paragraph
        return text


@registry.tool
def read_word(path: str, include_paragraphs: bool = True, include_tables: bool = True) -> str:
    """
    Read a Word document from the local filesystem and extract text content.

    Parameters
    -----
    path : str
        The path to the Word document, relative to the current working directory
    include_paragraphs : bool, optional
        Whether to include paragraph text (default: True)
    include_tables : bool, optional
        Whether to include table content (default: True)
    """
    if not path:
        raise ValueError("File path cannot be empty")

    abs_path = _resolve_path(path)

    if not abs_path.exists():
        raise FileNotFoundError(f"Word document not found: {path}")

    doc = Document(str(abs_path))

    # Extract basic document info
    properties = doc.core_properties
    doc_info = []
    if properties.title:
        doc_info.append(f"Title: {properties.title}")
    if properties.author:
        doc_info.append(f"Author: {properties.author}")
    if properties.subject:
        doc_info.append(f"Subject: {properties.subject}")

    content_parts = []
    
    # Add document info if available
    if doc_info:
        content_parts.append("Document Information:\n" + "\n".join(doc_info))

    # Extract paragraphs with markdown-like formatting
    if include_paragraphs:
        paragraphs = []
        for paragraph in doc.paragraphs:
            formatted_text = _format_paragraph_with_style(paragraph)
            if formatted_text:  # Only include non-empty paragraphs
                paragraphs.append(formatted_text)
        
        if paragraphs:
            content_parts.append("Content:\n" + "\n\n".join(paragraphs))

    # Extract tables
    if include_tables:
        for i, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Only include rows with content
                    table_rows.append(" | ".join(row_data))
            
            if table_rows:
                table_content = f"## Table {i + 1}\n\n" + "\n".join(table_rows)
                content_parts.append(table_content)

    return "\n\n".join(content_parts) if content_parts else "No content found in document."